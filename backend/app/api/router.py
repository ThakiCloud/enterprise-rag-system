from fastapi import APIRouter, Depends, UploadFile, File, HTTPException, Body
from fastapi.responses import StreamingResponse
from datetime import datetime
from typing import List, Dict, Any, AsyncGenerator
import logging
import requests
from bs4 import BeautifulSoup
import re
import json
import asyncio
import io
import os
from pathlib import Path

# Document processing imports
from docx import Document
import pypdf

from ..schemas.query import QueryRequest, QueryResponse
from ..schemas.document import DocumentUploadResponse
from ..schemas.session import SessionInfo, SessionMemoryRequest, UserMemory
from ..core.dependencies import get_rag_agent, get_research_team, get_knowledge_base, SimpleAgent, SimpleKnowledgeBase
from ..core.memory_manager import session_memory_manager

router = APIRouter()

# Set up logging
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = pypdf.PdfReader(pdf_file)
        
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        return text
    except Exception as e:
        logger.error(f"PDF processing failed: {e}")
        raise ValueError(f"PDF 파일을 읽을 수 없습니다: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        return text
    except Exception as e:
        logger.error(f"DOCX processing failed: {e}")
        raise ValueError(f"DOCX 파일을 읽을 수 없습니다: {str(e)}")

def extract_text_from_txt_or_md(file_content: bytes) -> str:
    """Extract text from TXT or MD file"""
    try:
        # Try different encodings
        encodings = ['utf-8', 'utf-8-sig', 'cp949', 'euc-kr', 'latin-1']
        
        for encoding in encodings:
            try:
                return file_content.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, use utf-8 with error handling
        return file_content.decode('utf-8', errors='ignore')
    except Exception as e:
        logger.error(f"Text file processing failed: {e}")
        raise ValueError(f"텍스트 파일을 읽을 수 없습니다: {str(e)}")

def process_document(file_content: bytes, filename: str, file_extension: str) -> Dict[str, Any]:
    """Process document and extract text content"""
    try:
        if file_extension.lower() == '.pdf':
            text = extract_text_from_pdf(file_content)
        elif file_extension.lower() == '.docx':
            text = extract_text_from_docx(file_content)
        elif file_extension.lower() in ['.txt', '.md']:
            text = extract_text_from_txt_or_md(file_content)
        else:
            raise ValueError(f"지원하지 않는 파일 형식입니다: {file_extension}")
        
        # Clean and validate text
        text = text.strip()
        if not text:
            raise ValueError("파일에서 텍스트를 추출할 수 없습니다. 파일이 비어있거나 손상되었을 수 있습니다.")
        
        # Calculate basic statistics
        word_count = len(text.split())
        char_count = len(text)
        line_count = len(text.splitlines())
        
        return {
            'text': text,
            'word_count': word_count,
            'char_count': char_count,
            'line_count': line_count,
            'filename': filename,
            'file_type': file_extension,
            'status': 'success'
        }
        
    except Exception as e:
        logger.error(f"Document processing failed for {filename}: {e}")
        return {
            'filename': filename,
            'file_type': file_extension,
            'error': str(e),
            'status': 'error'
        }

def extract_web_content(url: str) -> Dict[str, Any]:
    """Extract content from a web URL"""
    try:
        # Set headers to mimic a real browser
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        # Fetch the webpage
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        # Parse HTML content
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer", "header", "aside"]):
            script.decompose()
        
        # Extract title
        title = soup.find('title')
        title_text = title.get_text().strip() if title else "No title"
        
        # Extract main content
        # Try to find main content areas
        main_content = soup.find('main') or soup.find('article') or soup.find('div', class_=re.compile(r'content|main|article', re.I))
        
        if main_content:
            content_text = main_content.get_text()
        else:
            # Fallback to body content
            content_text = soup.get_text()
        
        # Clean up the text
        lines = (line.strip() for line in content_text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        content_text = ' '.join(chunk for chunk in chunks if chunk)
        
        # Extract meta description
        meta_desc = soup.find('meta', attrs={'name': 'description'})
        description = meta_desc.get('content', '') if meta_desc else ''
        
        # Extract headings for structure
        headings = []
        for heading in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            headings.append(f"{heading.name}: {heading.get_text().strip()}")
        
        return {
            'url': url,
            'title': title_text,
            'description': description,
            'content': content_text[:5000],  # Limit content length
            'headings': headings[:10],  # Limit number of headings
            'word_count': len(content_text.split()),
            'status': 'success'
        }
        
    except requests.RequestException as e:
        logger.error(f"Error fetching URL {url}: {e}")
        return {
            'url': url,
            'error': f"Failed to fetch URL: {str(e)}",
            'status': 'error'
        }
    except Exception as e:
        logger.error(f"Error processing URL {url}: {e}")
        return {
            'url': url,
            'error': f"Failed to process content: {str(e)}",
            'status': 'error'
        }

@router.post("/query/", response_model=QueryResponse)
async def query_knowledge(
    request: QueryRequest
):
    """Process a query using the RAG system."""
    try:
        logger.info(f"Received query request for session: {request.session_id}")
        session_id = request.session_id or f"session_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        logger.info(f"Getting agent for query (advanced_reasoning={request.use_advanced_reasoning})...")
        rag_agent = get_rag_agent()
        research_team = get_research_team()
        agent = research_team if request.use_advanced_reasoning else rag_agent
        logger.info(f"Using agent: {agent.name if hasattr(agent, 'name') else 'SimpleAgent'}")
        
        logger.info(f"Executing agent with question: '{request.question}'")
        response = await agent.arun(request.question)
        logger.info("Agent execution finished, creating response.")
        
        # Simple response handling
        sources = []
        reasoning_steps = None
        
        return QueryResponse(
            answer=response,
            session_id=session_id,
            user_id=request.user_id,
            sources=sources,
            processing_time=None,
            memory_updated=False,
            memory_count=0,
            tokens_used=None,
            model_used=None
        )
    except Exception as e:
        logger.error(f"Error processing query: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error processing query: {str(e)}")


@router.post("/query/stream/")
async def query_knowledge_stream(
    request: QueryRequest
):
    """Process a query using the RAG system with streaming response and memory support."""
    try:
        logger.info(f"Received streaming query request for session: {request.session_id}")
        session_id = request.session_id or f"session_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        user_id = request.user_id or session_id  # Use session_id as fallback user_id
        
        logger.info(f"Getting agent for streaming query (advanced_reasoning={request.use_advanced_reasoning})...")
        rag_agent = get_rag_agent()
        research_team = get_research_team()
        agent = research_team if request.use_advanced_reasoning else rag_agent
        logger.info(f"Using agent: {agent.name if hasattr(agent, 'name') else 'SimpleAgent'}")
        
        # Get relevant memories for context if memory is enabled
        relevant_memories = []
        if request.use_memory and user_id:
            try:
                relevant_memories = session_memory_manager.get_relevant_memories_for_query(
                    user_id=user_id, 
                    query=request.question,
                    limit=3
                )
                logger.info(f"Found {len(relevant_memories)} relevant memories for user {user_id}")
            except Exception as e:
                logger.warning(f"Failed to get relevant memories: {e}")
        
        async def generate_response():
            try:
                logger.info(f"Executing agent with question: '{request.question}'")
                
                # Prepare context with relevant memories
                context_info = ""
                if relevant_memories:
                    memory_context = "\n".join([
                        f"- {mem.memory} (관련 토픽: {', '.join(mem.topics)})" 
                        for mem in relevant_memories
                    ])
                    context_info = f"\n\n이전 대화 내용:\n{memory_context}\n"
                
                # Add context to question if memories exist
                enhanced_question = request.question
                if context_info:
                    enhanced_question = f"{request.question}{context_info}"
                
                response = await agent.arun(enhanced_question, user_id=user_id, session_id=session_id)
                logger.info("Agent execution finished, starting streaming response.")
                
                # Parse response to handle <think> tags
                think_content = ""
                main_response = response
                
                # Extract <think> content if present
                if "<think>" in response and "</think>" in response:
                    think_start = response.find("<think>")
                    think_end = response.find("</think>") + 8
                    think_content = response[think_start:think_end]
                    main_response = response[think_end:].strip()
                
                # Create response data
                response_data = {
                    "query": request.question,
                    "session_id": session_id,
                    "timestamp": datetime.now().isoformat(),
                    "status": "streaming"
                }
                
                # Send initial response metadata
                yield f"data: {json.dumps(response_data)}\n\n"
                
                # Stream the response word by word
                words = main_response.split()
                total_words = len(words)
                
                for i, word in enumerate(words):
                    # Add thinking section for first word if present
                    if i == 0 and think_content:
                        word_data = {
                            "word": word,
                            "index": i,
                            "total": total_words,
                            "progress": round((i + 1) / total_words * 100, 1),
                            "think": think_content,
                            "is_complete": False
                        }
                    else:
                        word_data = {
                            "word": word,
                            "index": i,
                            "total": total_words,
                            "progress": round((i + 1) / total_words * 100, 1),
                            "is_complete": False
                        }
                    
                    yield f"data: {json.dumps(word_data)}\n\n"
                    
                    # Add delay for streaming effect
                    if word.endswith(('.', '!', '?', ':')):
                        await asyncio.sleep(0.3)  # Longer pause for sentence endings
                    elif word.endswith(','):
                        await asyncio.sleep(0.15)  # Medium pause for commas
                    else:
                        await asyncio.sleep(0.05)  # Short pause for regular words
                
                # Create memory from conversation if enabled
                memory_updated = False
                memory_count = 0
                if request.use_memory and user_id:
                    try:
                        # Create memories from this conversation
                        conversation_messages = [
                            {"role": "user", "content": request.question},
                            {"role": "assistant", "content": main_response}
                        ]
                        session_memory_manager.create_memories_from_conversation(
                            user_id=user_id,
                            messages=conversation_messages
                        )
                        memory_updated = True
                        memory_count = session_memory_manager.get_memory_count(user_id)
                        logger.info(f"Updated memories for user {user_id}, total count: {memory_count}")
                    except Exception as e:
                        logger.warning(f"Failed to create memories: {e}")
                
                # Send completion signal
                completion_data = {
                    "is_complete": True,
                    "total_words": total_words,
                    "full_response": main_response,
                    "think_content": think_content if think_content else None,
                    "session_id": session_id,
                    "user_id": user_id,
                    "memory_updated": memory_updated,
                    "memory_count": memory_count,
                    "relevant_memories_count": len(relevant_memories),
                    "timestamp": datetime.now().isoformat(),
                    "status": "completed"
                }
                yield f"data: {json.dumps(completion_data)}\n\n"
                
            except Exception as e:
                logger.error(f"Error in streaming response: {str(e)}", exc_info=True)
                error_data = {
                    "error": str(e),
                    "is_complete": True,
                    "status": "error"
                }
                yield f"data: {json.dumps(error_data)}\n\n"
        
        return StreamingResponse(
            generate_response(),
            media_type="text/plain",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "Content-Type": "text/event-stream",
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
                "Access-Control-Allow-Headers": "Content-Type"
            }
        )
        
    except Exception as e:
        logger.error(f"Error setting up streaming query: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error setting up streaming query: {str(e)}")


@router.post("/upload-document/", response_model=DocumentUploadResponse)
async def upload_document(
    file: UploadFile = File(...)
):
    """Upload and process a document into the knowledge base."""
    try:
        # Validate file type
        allowed_extensions = ['.pdf', '.docx', '.txt', '.md']
        file_extension = '.' + file.filename.lower().split('.')[-1] if '.' in file.filename else ''
        
        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type: {file_extension}. Allowed types: {', '.join(allowed_extensions)}"
            )
        
        # Validate file size (10MB limit)
        max_size = 10 * 1024 * 1024  # 10MB
        file_content = await file.read()
        if len(file_content) > max_size:
            raise HTTPException(status_code=400, detail="File size exceeds 10MB limit")
        
        # Process document
        document_data = process_document(file_content, file.filename, file_extension)
        
        if document_data['status'] == 'error':
            raise HTTPException(status_code=400, detail=document_data['error'])
        
        # Add to knowledge base
        knowledge_base = get_knowledge_base()
        metadata = {
            "filename": file.filename,
            "type": file_extension,
            "word_count": document_data['word_count'],
            "char_count": document_data['char_count'],
            "line_count": document_data['line_count'],
            "extracted_at": datetime.now().isoformat()
        }
        
        knowledge_base.add_document(document_data['text'], metadata)
        
        logger.info(f"Successfully processed and added document: {file.filename}")
        
        return DocumentUploadResponse(
            message=f"문서 '{file.filename}'이 성공적으로 업로드되고 분석되었습니다.",
            document_id=f"doc_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{len(knowledge_base.documents)}",
            filename=file.filename,
            status="success",
            metadata={
                "word_count": document_data['word_count'],
                "char_count": document_data['char_count'],
                "line_count": document_data['line_count'],
                "file_type": file_extension
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing document {file.filename}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"문서 처리 중 오류가 발생했습니다: {str(e)}")


@router.post("/add-url/")
async def add_url_endpoint(
    url: str = Body(..., embed=True)
):
    """Add URL content to the knowledge base."""
    try:
        logger.info(f"Processing URL: {url}")
        
        # Basic URL validation
        if not url.startswith(('http://', 'https://')):
            raise HTTPException(status_code=400, detail="URL must start with http:// or https://")
        
        # Extract web content
        content_data = extract_web_content(url)
        
        if content_data['status'] == 'error':
            raise HTTPException(status_code=400, detail=content_data['error'])
        
        # Prepare content for knowledge base
        formatted_content = f"""
URL: {content_data['url']}
Title: {content_data['title']}
Description: {content_data['description']}

Headings:
{chr(10).join(content_data['headings'])}

Content:
{content_data['content']}
        """.strip()
        
        # Add to knowledge base
        knowledge_base = get_knowledge_base()
        metadata = {
            "url": url,
            "title": content_data['title'],
            "type": "web_content",
            "word_count": content_data['word_count'],
            "extracted_at": datetime.now().isoformat()
        }
        
        knowledge_base.add_document(formatted_content, metadata)
        
        logger.info(f"Successfully added URL content to knowledge base: {url}")
        
        return {
            "url": url,
            "title": content_data['title'],
            "word_count": content_data['word_count'],
            "status": "success",
            "message": f"Successfully extracted and added content from '{content_data['title']}'"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing URL {url}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error processing URL: {str(e)}")


@router.post("/analyze-url/")
async def analyze_url_endpoint(
    url: str = Body(..., embed=True),
    question: str = Body(default="Analyze this web content and provide a comprehensive summary.", embed=True)
):
    """Extract content from URL and immediately analyze it."""
    try:
        logger.info(f"Analyzing URL: {url}")
        
        # First, add the URL to knowledge base
        url_result = await add_url_endpoint(url)
        
        if url_result['status'] != 'success':
            raise HTTPException(status_code=400, detail="Failed to extract URL content")
        
        # Now analyze the content
        rag_agent = get_rag_agent()
        
        # Create analysis question that references the URL
        analysis_question = f"Based on the content from {url} (titled '{url_result['title']}'), {question}"
        
        logger.info(f"Analyzing with question: {analysis_question}")
        response = await rag_agent.arun(analysis_question)
        
        return {
            "url": url,
            "title": url_result['title'],
            "word_count": url_result['word_count'],
            "question": question,
            "analysis": response,
            "status": "success",
            "timestamp": datetime.now().isoformat()
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error analyzing URL {url}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error analyzing URL: {str(e)}")


@router.get("/sessions/", response_model=List[SessionInfo])
async def list_sessions():
    """List all active sessions."""
    try:
        # Simple implementation
        sessions = [SessionInfo(
            session_id="example_session",
            created_at=datetime.now().isoformat(),
            last_activity=datetime.now().isoformat(),
            query_count=0
        )]
        
        return sessions
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error listing sessions: {str(e)}")


@router.delete("/sessions/{session_id}")
async def delete_session(
    session_id: str
):
    """Delete a specific session."""
    try:
        return {
            "message": f"Session {session_id} deleted successfully",
            "session_id": session_id,
            "status": "success"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting session: {str(e)}")


@router.get("/sessions/{session_id}")
async def get_session(
    session_id: str
):
    """Get information about a specific session."""
    try:
        # Use session_id as user_id for memory lookup
        user_id = session_id
        memory_count = session_memory_manager.get_memory_count(user_id)
        
        return SessionInfo(
            session_id=session_id,
            user_id=user_id,
            created_at=datetime.now().isoformat(),
            last_activity=datetime.now().isoformat(),
            query_count=0,
            has_memory=memory_count > 0,
            memory_count=memory_count
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error getting session info: {str(e)}")


@router.post("/sessions/{session_id}/memories")
async def manage_session_memory(
    session_id: str,
    request: SessionMemoryRequest
):
    """Manage memories for a session"""
    try:
        user_id = request.user_id or session_id
        
        if request.action == "get":
            memories = session_memory_manager.get_user_memories(user_id, limit=20)
            return {
                "session_id": session_id,
                "user_id": user_id,
                "memories": memories,
                "total_count": len(memories),
                "status": "success"
            }
        
        elif request.action == "add":
            if not request.memory_content:
                raise HTTPException(status_code=400, detail="Memory content is required for add action")
            
            memory_id = session_memory_manager.add_user_memory(
                user_id=user_id,
                memory_content=request.memory_content,
                topics=request.topics or []
            )
            
            return {
                "session_id": session_id,
                "user_id": user_id,
                "memory_id": memory_id,
                "message": "Memory added successfully",
                "status": "success"
            }
        
        elif request.action == "delete":
            if not request.memory_id:
                raise HTTPException(status_code=400, detail="Memory ID is required for delete action")
            
            success = session_memory_manager.delete_user_memory(user_id, request.memory_id)
            
            if success:
                return {
                    "session_id": session_id,
                    "user_id": user_id,
                    "message": "Memory deleted successfully",
                    "status": "success"
                }
            else:
                raise HTTPException(status_code=404, detail="Memory not found or could not be deleted")
        
        elif request.action == "clear":
            success = session_memory_manager.clear_user_memories(user_id)
            
            if success:
                return {
                    "session_id": session_id,
                    "user_id": user_id,
                    "message": "All memories cleared successfully",
                    "status": "success"
                }
            else:
                raise HTTPException(status_code=500, detail="Failed to clear memories")
        
        else:
            raise HTTPException(status_code=400, detail="Invalid action. Use 'get', 'add', 'delete', or 'clear'")
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error managing session memory: {e}")
        raise HTTPException(status_code=500, detail=f"Error managing session memory: {str(e)}")


@router.get("/sessions/{session_id}/memories")
async def get_session_memories(
    session_id: str,
    limit: int = 10
):
    """Get memories for a session"""
    try:
        user_id = session_id  # Use session_id as user_id
        memories = session_memory_manager.get_user_memories(user_id, limit=limit)
        
        return {
            "session_id": session_id,
            "user_id": user_id,
            "memories": memories,
            "total_count": len(memories),
            "status": "success"
        }
    except Exception as e:
        logger.error(f"Error getting session memories: {e}")
        raise HTTPException(status_code=500, detail=f"Error getting session memories: {str(e)}")


@router.post("/sessions/{session_id}/search-memories")
async def search_session_memories(
    session_id: str,
    query: str = Body(..., embed=True),
    limit: int = Body(5, embed=True)
):
    """Search memories for a session based on query"""
    try:
        user_id = session_id  # Use session_id as user_id
        memories = session_memory_manager.search_user_memories(user_id, query, limit)
        
        return {
            "session_id": session_id,
            "user_id": user_id,
            "query": query,
            "memories": memories,
            "found_count": len(memories),
            "status": "success"
        }
    except Exception as e:
        logger.error(f"Error searching session memories: {e}")
        raise HTTPException(status_code=500, detail=f"Error searching session memories: {str(e)}")


@router.get("/knowledge-base/stats")
async def get_knowledge_base_stats():
    """Get statistics about the knowledge base."""
    try:
        knowledge_base = get_knowledge_base()
        
        # Count different types of documents
        total_docs = len(knowledge_base.documents)
        web_docs = sum(1 for doc in knowledge_base.documents if doc.get('metadata', {}).get('type') == 'web_content')
        file_docs = total_docs - web_docs
        
        stats = {
            "total_documents": total_docs,
            "web_documents": web_docs,
            "file_documents": file_docs,
            "status": "active"
        }
        
        return stats
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error getting knowledge base stats: {str(e)}")


@router.post("/analyze-document/")
async def analyze_document(
    file: UploadFile = File(...),
    question: str = Body(default="이 문서의 내용을 분석하고 주요 내용을 요약해주세요.", embed=True)
):
    """Upload and immediately analyze a document."""
    try:
        logger.info(f"Analyzing document: {file.filename}")
        
        # First upload the document
        upload_result = await upload_document(file)
        
        if upload_result.status != "success":
            raise HTTPException(status_code=400, detail="문서 업로드에 실패했습니다.")
        
        # Now analyze the document
        rag_agent = get_rag_agent()
        
        # Create analysis question that references the document
        analysis_question = f"방금 업로드된 문서 '{upload_result.filename}'에 대해: {question}"
        
        logger.info(f"Analyzing document with question: {analysis_question}")
        response = await rag_agent.arun(analysis_question)
        
        return {
            "filename": upload_result.filename,
            "document_id": upload_result.document_id,
            "metadata": upload_result.metadata,
            "question": question,
            "analysis": response,
            "status": "success",
            "timestamp": datetime.now().isoformat()
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error analyzing document {file.filename}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"문서 분석 중 오류가 발생했습니다: {str(e)}")


@router.post("/upload-multiple-documents/")
async def upload_multiple_documents(
    files: List[UploadFile] = File(...)
):
    """Upload multiple documents at once."""
    try:
        if len(files) > 10:
            raise HTTPException(status_code=400, detail="한 번에 최대 10개의 파일만 업로드할 수 있습니다.")
        
        results = []
        success_count = 0
        error_count = 0
        
        for file in files:
            try:
                # Process each file
                file_content = await file.read()
                
                # Reset file position for reuse
                await file.seek(0)
                
                # Validate file
                allowed_extensions = ['.pdf', '.docx', '.txt', '.md']
                file_extension = '.' + file.filename.lower().split('.')[-1] if '.' in file.filename else ''
                
                if file_extension not in allowed_extensions:
                    results.append({
                        "filename": file.filename,
                        "status": "error",
                        "error": f"지원하지 않는 파일 형식: {file_extension}"
                    })
                    error_count += 1
                    continue
                
                # Check file size
                max_size = 10 * 1024 * 1024  # 10MB
                if len(file_content) > max_size:
                    results.append({
                        "filename": file.filename,
                        "status": "error",
                        "error": "파일 크기가 10MB를 초과합니다."
                    })
                    error_count += 1
                    continue
                
                # Process document
                document_data = process_document(file_content, file.filename, file_extension)
                
                if document_data['status'] == 'error':
                    results.append({
                        "filename": file.filename,
                        "status": "error",
                        "error": document_data['error']
                    })
                    error_count += 1
                    continue
                
                # Add to knowledge base
                knowledge_base = get_knowledge_base()
                metadata = {
                    "filename": file.filename,
                    "type": file_extension,
                    "word_count": document_data['word_count'],
                    "char_count": document_data['char_count'],
                    "line_count": document_data['line_count'],
                    "extracted_at": datetime.now().isoformat()
                }
                
                knowledge_base.add_document(document_data['text'], metadata)
                
                results.append({
                    "filename": file.filename,
                    "status": "success",
                    "document_id": f"doc_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{len(knowledge_base.documents)}",
                    "metadata": {
                        "word_count": document_data['word_count'],
                        "char_count": document_data['char_count'],
                        "line_count": document_data['line_count'],
                        "file_type": file_extension
                    }
                })
                success_count += 1
                
            except Exception as e:
                results.append({
                    "filename": file.filename,
                    "status": "error",
                    "error": str(e)
                })
                error_count += 1
        
        return {
            "message": f"총 {len(files)}개 파일 처리 완료: {success_count}개 성공, {error_count}개 실패",
            "total_files": len(files),
            "success_count": success_count,
            "error_count": error_count,
            "results": results,
            "status": "completed",
            "timestamp": datetime.now().isoformat()
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing multiple documents: {str(e)}")
        raise HTTPException(status_code=500, detail=f"다중 문서 업로드 중 오류가 발생했습니다: {str(e)}") 